#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Export Manager Module
Handles exporting documents to various formats (PDF, DOCX)
"""

from docx import Document
from docx.shared import Pt, RGBColor
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.pdfbase import pdfmetrics
from reportlab.lib.enums import TA_JUSTIFY, TA_LEFT
import os


class ExportManager:
    """Manages document export operations"""
    
    def __init__(self):
        """Initialize the export manager"""
        pass
    
    def export_to_pdf(self, text, file_path):
        """
        Export text to PDF format
        
        Args:
            text (str): Text content to export
            file_path (str): Output file path
        """
        try:
            # Create PDF document
            doc = SimpleDocTemplate(
                file_path,
                pagesize=A4,
                rightMargin=72,
                leftMargin=72,
                topMargin=72,
                bottomMargin=18,
            )
            
            # Container for the 'Flowable' objects
            elements = []
            
            # Define styles
            styles = getSampleStyleSheet()
            
            # Create custom style for Turkish text
            custom_style = ParagraphStyle(
                'CustomStyle',
                parent=styles['Normal'],
                fontName='Helvetica',
                fontSize=12,
                leading=16,
                alignment=TA_LEFT,
                spaceAfter=12,
            )
            
            # Split text into paragraphs
            paragraphs = text.split('\n')
            
            for para_text in paragraphs:
                if para_text.strip():
                    # Create paragraph
                    para = Paragraph(para_text, custom_style)
                    elements.append(para)
                    elements.append(Spacer(1, 0.2*inch))
            
            # Build PDF
            doc.build(elements)
            
        except Exception as e:
            raise Exception(f"PDF export error: {str(e)}")
    
    def export_to_docx(self, text, file_path):
        """
        Export text to DOCX format
        
        Args:
            text (str): Text content to export
            file_path (str): Output file path
        """
        try:
            # Create a new Document
            document = Document()
            
            # Set default font
            style = document.styles['Normal']
            font = style.font
            font.name = 'Arial'
            font.size = Pt(12)
            
            # Add title
            document.add_heading('Gerekçeli Karar', 0)
            
            # Split text into paragraphs and add to document
            paragraphs = text.split('\n')
            
            for para_text in paragraphs:
                if para_text.strip():
                    para = document.add_paragraph(para_text)
                    # Set paragraph formatting
                    para_format = para.paragraph_format
                    para_format.line_spacing = 1.5
                    para_format.space_after = Pt(12)
                else:
                    # Add empty paragraph for spacing
                    document.add_paragraph()
            
            # Save the document
            document.save(file_path)
            
        except Exception as e:
            raise Exception(f"DOCX export error: {str(e)}")
    
    def export_to_txt(self, text, file_path):
        """
        Export text to plain text format
        
        Args:
            text (str): Text content to export
            file_path (str): Output file path
        """
        try:
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(text)
        except Exception as e:
            raise Exception(f"TXT export error: {str(e)}")
    
    def export_with_formatting(self, text, file_path, format_info=None):
        """
        Export text with formatting information
        
        Args:
            text (str): Text content to export
            file_path (str): Output file path
            format_info (dict): Formatting information (font, size, color, etc.)
        """
        # Determine format from file extension
        _, ext = os.path.splitext(file_path)
        ext = ext.lower()
        
        if ext == '.pdf':
            self.export_to_pdf(text, file_path)
        elif ext == '.docx':
            self.export_to_docx(text, file_path)
        elif ext == '.txt':
            self.export_to_txt(text, file_path)
        else:
            raise Exception(f"Unsupported file format: {ext}")
