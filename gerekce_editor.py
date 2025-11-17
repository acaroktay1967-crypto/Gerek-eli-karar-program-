#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Gerekçeli Karar Programı
Turkish Legal Document Editor with Translation and Speech Recognition
"""

import tkinter as tk
from tkinter import ttk, scrolledtext, messagebox, filedialog, font as tkfont
import threading
import os

# Import for translation
try:
    from googletrans import Translator
except ImportError:
    Translator = None

# Import for speech recognition
try:
    import speech_recognition as sr
except ImportError:
    sr = None

# Import for document export
try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
except ImportError:
    Document = None

# Import for PDF export
try:
    from fpdf import FPDF
except ImportError:
    FPDF = None


class GerekceliKararProgram:
    """Main application class for the Legal Document Editor"""
    
    def __init__(self, root):
        self.root = root
        self.root.title("Gerekçeli Karar Programı")
        self.root.geometry("1000x700")
        
        # Initialize translator
        self.translator = Translator() if Translator else None
        
        # Initialize speech recognizer
        self.recognizer = sr.Recognizer() if sr else None
        
        # Text formatting state
        self.is_bold = False
        self.is_italic = False
        self.is_underline = False
        self.current_font_family = "Arial"
        self.current_font_size = 12
        
        # Create UI
        self.create_menu()
        self.create_toolbar()
        self.create_text_area()
        self.create_statusbar()
        
        # Configure tags for text formatting
        self.configure_text_tags()
        
    def create_menu(self):
        """Create menu bar"""
        menubar = tk.Menu(self.root)
        self.root.config(menu=menubar)
        
        # File menu
        file_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="Dosya", menu=file_menu)
        file_menu.add_command(label="Yeni", command=self.new_document)
        file_menu.add_command(label="Aç", command=self.open_document)
        file_menu.add_command(label="Kaydet", command=self.save_document)
        file_menu.add_separator()
        file_menu.add_command(label="PDF Olarak Kaydet", command=self.export_to_pdf)
        file_menu.add_command(label="Word Olarak Kaydet", command=self.export_to_docx)
        file_menu.add_separator()
        file_menu.add_command(label="Çıkış", command=self.root.quit)
        
        # Edit menu
        edit_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="Düzenle", menu=edit_menu)
        edit_menu.add_command(label="Geri Al", command=self.undo)
        edit_menu.add_command(label="Yinele", command=self.redo)
        edit_menu.add_separator()
        edit_menu.add_command(label="Kes", command=self.cut)
        edit_menu.add_command(label="Kopyala", command=self.copy)
        edit_menu.add_command(label="Yapıştır", command=self.paste)
        
        # Translation menu
        translate_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="Çeviri", menu=translate_menu)
        translate_menu.add_command(label="İngilizce'ye Çevir", 
                                  command=lambda: self.translate_text('en'))
        translate_menu.add_command(label="Almanca'ya Çevir", 
                                  command=lambda: self.translate_text('de'))
        translate_menu.add_command(label="İspanyolca'ya Çevir", 
                                  command=lambda: self.translate_text('es'))
        translate_menu.add_command(label="Fransızca'ya Çevir", 
                                  command=lambda: self.translate_text('fr'))
        
        # Tools menu
        tools_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="Araçlar", menu=tools_menu)
        tools_menu.add_command(label="Sesle Metin Ekle", command=self.start_dictation)
        tools_menu.add_command(label="Satıra İşaretleyici Ekle", command=self.add_line_marker)
        
        # Help menu
        help_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="Yardım", menu=help_menu)
        help_menu.add_command(label="Hakkında", command=self.show_about)
        
    def create_toolbar(self):
        """Create toolbar with formatting options"""
        toolbar = tk.Frame(self.root, relief=tk.RAISED, borderwidth=2)
        toolbar.pack(side=tk.TOP, fill=tk.X, padx=5, pady=5)
        
        # Font family selector
        tk.Label(toolbar, text="Font:").pack(side=tk.LEFT, padx=5)
        self.font_family_var = tk.StringVar(value=self.current_font_family)
        font_families = ["Arial", "Times New Roman", "Courier New", "Verdana", "Georgia"]
        font_dropdown = ttk.Combobox(toolbar, textvariable=self.font_family_var, 
                                     values=font_families, width=15, state='readonly')
        font_dropdown.pack(side=tk.LEFT, padx=5)
        font_dropdown.bind('<<ComboboxSelected>>', self.change_font_family)
        
        # Font size selector
        tk.Label(toolbar, text="Boyut:").pack(side=tk.LEFT, padx=5)
        self.font_size_var = tk.IntVar(value=self.current_font_size)
        font_sizes = [8, 9, 10, 11, 12, 14, 16, 18, 20, 22, 24, 26, 28, 36, 48, 72]
        size_dropdown = ttk.Combobox(toolbar, textvariable=self.font_size_var, 
                                     values=font_sizes, width=5, state='readonly')
        size_dropdown.pack(side=tk.LEFT, padx=5)
        size_dropdown.bind('<<ComboboxSelected>>', self.change_font_size)
        
        # Bold button
        self.bold_btn = tk.Button(toolbar, text="B", font=("Arial", 10, "bold"),
                                 width=3, command=self.toggle_bold)
        self.bold_btn.pack(side=tk.LEFT, padx=2)
        
        # Italic button
        self.italic_btn = tk.Button(toolbar, text="I", font=("Arial", 10, "italic"),
                                   width=3, command=self.toggle_italic)
        self.italic_btn.pack(side=tk.LEFT, padx=2)
        
        # Underline button
        self.underline_btn = tk.Button(toolbar, text="U", font=("Arial", 10, "underline"),
                                      width=3, command=self.toggle_underline)
        self.underline_btn.pack(side=tk.LEFT, padx=2)
        
        # Separator
        tk.Frame(toolbar, width=2, relief=tk.SUNKEN, borderwidth=1).pack(side=tk.LEFT, 
                                                                          fill=tk.Y, padx=10)
        
        # Dictation button
        self.dictation_btn = tk.Button(toolbar, text="🎤 Dikte", 
                                      command=self.start_dictation)
        self.dictation_btn.pack(side=tk.LEFT, padx=5)
        
        # Line marker button
        tk.Button(toolbar, text="📌 İşaretle", 
                 command=self.add_line_marker).pack(side=tk.LEFT, padx=5)
        
    def create_text_area(self):
        """Create main text editing area"""
        # Create frame for text area
        text_frame = tk.Frame(self.root)
        text_frame.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        
        # Create text widget with scrollbar
        self.text_area = tk.Text(text_frame, wrap=tk.WORD, undo=True, 
                                font=(self.current_font_family, self.current_font_size))
        scrollbar = tk.Scrollbar(text_frame, command=self.text_area.yview)
        self.text_area.configure(yscrollcommand=scrollbar.set)
        
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        self.text_area.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        
        # Bind events
        self.text_area.bind('<KeyRelease>', self.update_statusbar)
        self.text_area.bind('<ButtonRelease-1>', self.update_statusbar)
        
    def create_statusbar(self):
        """Create status bar"""
        self.statusbar = tk.Label(self.root, text="Hazır", anchor=tk.W, 
                                 relief=tk.SUNKEN, borderwidth=1)
        self.statusbar.pack(side=tk.BOTTOM, fill=tk.X)
        
    def configure_text_tags(self):
        """Configure tags for text formatting"""
        # Bold tag
        bold_font = tkfont.Font(self.text_area, self.text_area.cget("font"))
        bold_font.configure(weight="bold")
        self.text_area.tag_configure("bold", font=bold_font)
        
        # Italic tag
        italic_font = tkfont.Font(self.text_area, self.text_area.cget("font"))
        italic_font.configure(slant="italic")
        self.text_area.tag_configure("italic", font=italic_font)
        
        # Underline tag
        self.text_area.tag_configure("underline", underline=True)
        
        # Bold+Italic tag
        bold_italic_font = tkfont.Font(self.text_area, self.text_area.cget("font"))
        bold_italic_font.configure(weight="bold", slant="italic")
        self.text_area.tag_configure("bold_italic", font=bold_italic_font)
        
        # Line marker tag
        self.text_area.tag_configure("marker", background="yellow")
        
    def toggle_bold(self):
        """Toggle bold formatting"""
        self.is_bold = not self.is_bold
        self.bold_btn.config(relief=tk.SUNKEN if self.is_bold else tk.RAISED)
        self.apply_formatting()
        
    def toggle_italic(self):
        """Toggle italic formatting"""
        self.is_italic = not self.is_italic
        self.italic_btn.config(relief=tk.SUNKEN if self.is_italic else tk.RAISED)
        self.apply_formatting()
        
    def toggle_underline(self):
        """Toggle underline formatting"""
        self.is_underline = not self.is_underline
        self.underline_btn.config(relief=tk.SUNKEN if self.is_underline else tk.RAISED)
        self.apply_formatting()
        
    def apply_formatting(self):
        """Apply current formatting to selected text"""
        try:
            # Get selected text range
            start = self.text_area.index(tk.SEL_FIRST)
            end = self.text_area.index(tk.SEL_LAST)
            
            # Remove all formatting tags
            for tag in ["bold", "italic", "underline", "bold_italic"]:
                self.text_area.tag_remove(tag, start, end)
            
            # Apply formatting based on current state
            if self.is_bold and self.is_italic:
                self.text_area.tag_add("bold_italic", start, end)
            elif self.is_bold:
                self.text_area.tag_add("bold", start, end)
            elif self.is_italic:
                self.text_area.tag_add("italic", start, end)
                
            if self.is_underline:
                self.text_area.tag_add("underline", start, end)
                
        except tk.TclError:
            # No selection
            pass
            
    def change_font_family(self, event=None):
        """Change font family"""
        self.current_font_family = self.font_family_var.get()
        current_font = tkfont.Font(font=self.text_area['font'])
        new_font = tkfont.Font(family=self.current_font_family, 
                              size=self.current_font_size)
        self.text_area.configure(font=new_font)
        self.configure_text_tags()
        
    def change_font_size(self, event=None):
        """Change font size"""
        self.current_font_size = self.font_size_var.get()
        current_font = tkfont.Font(font=self.text_area['font'])
        new_font = tkfont.Font(family=self.current_font_family, 
                              size=self.current_font_size)
        self.text_area.configure(font=new_font)
        self.configure_text_tags()
        
    def new_document(self):
        """Create a new document"""
        if messagebox.askyesno("Yeni Belge", "Mevcut belgeyi temizlemek istiyor musunuz?"):
            self.text_area.delete(1.0, tk.END)
            self.statusbar.config(text="Yeni belge oluşturuldu")
            
    def open_document(self):
        """Open a text document"""
        filename = filedialog.askopenfilename(
            title="Belge Aç",
            filetypes=[("Metin Dosyaları", "*.txt"), ("Tüm Dosyalar", "*.*")]
        )
        if filename:
            try:
                with open(filename, 'r', encoding='utf-8') as file:
                    content = file.read()
                    self.text_area.delete(1.0, tk.END)
                    self.text_area.insert(1.0, content)
                    self.statusbar.config(text=f"Dosya açıldı: {filename}")
            except Exception as e:
                messagebox.showerror("Hata", f"Dosya açılırken hata: {str(e)}")
                
    def save_document(self):
        """Save document as text file"""
        filename = filedialog.asksaveasfilename(
            title="Belge Kaydet",
            defaultextension=".txt",
            filetypes=[("Metin Dosyaları", "*.txt"), ("Tüm Dosyalar", "*.*")]
        )
        if filename:
            try:
                content = self.text_area.get(1.0, tk.END)
                with open(filename, 'w', encoding='utf-8') as file:
                    file.write(content)
                self.statusbar.config(text=f"Dosya kaydedildi: {filename}")
            except Exception as e:
                messagebox.showerror("Hata", f"Dosya kaydedilirken hata: {str(e)}")
                
    def export_to_pdf(self):
        """Export document to PDF"""
        if not FPDF:
            messagebox.showerror("Hata", "PDF desteği yüklü değil. fpdf kütüphanesi gerekli.")
            return
            
        filename = filedialog.asksaveasfilename(
            title="PDF Olarak Kaydet",
            defaultextension=".pdf",
            filetypes=[("PDF Dosyaları", "*.pdf")]
        )
        if filename:
            try:
                pdf = FPDF()
                pdf.add_page()
                
                # Add Unicode font support for Turkish characters
                pdf.add_font('DejaVu', '', '/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf', uni=True)
                pdf.set_font('DejaVu', '', 12)
                
                # Get text content
                content = self.text_area.get(1.0, tk.END)
                
                # Split into lines and add to PDF
                for line in content.split('\n'):
                    if line.strip():
                        pdf.multi_cell(0, 10, line)
                    else:
                        pdf.ln()
                
                pdf.output(filename)
                self.statusbar.config(text=f"PDF kaydedildi: {filename}")
                messagebox.showinfo("Başarılı", "Belge PDF olarak kaydedildi!")
            except Exception as e:
                messagebox.showerror("Hata", f"PDF oluşturulurken hata: {str(e)}")
                
    def export_to_docx(self):
        """Export document to Word DOCX"""
        if not Document:
            messagebox.showerror("Hata", "Word desteği yüklü değil. python-docx kütüphanesi gerekli.")
            return
            
        filename = filedialog.asksaveasfilename(
            title="Word Olarak Kaydet",
            defaultextension=".docx",
            filetypes=[("Word Dosyaları", "*.docx")]
        )
        if filename:
            try:
                doc = Document()
                
                # Get text content
                content = self.text_area.get(1.0, tk.END)
                
                # Add paragraphs
                for line in content.split('\n'):
                    if line.strip():
                        paragraph = doc.add_paragraph(line)
                    else:
                        doc.add_paragraph()
                
                doc.save(filename)
                self.statusbar.config(text=f"Word belgesi kaydedildi: {filename}")
                messagebox.showinfo("Başarılı", "Belge Word olarak kaydedildi!")
            except Exception as e:
                messagebox.showerror("Hata", f"Word belgesi oluşturulurken hata: {str(e)}")
                
    def translate_text(self, target_lang):
        """Translate selected text or entire document"""
        if not self.translator:
            messagebox.showerror("Hata", "Çeviri desteği yüklü değil. googletrans kütüphanesi gerekli.")
            return
            
        try:
            # Try to get selected text first
            try:
                text_to_translate = self.text_area.get(tk.SEL_FIRST, tk.SEL_LAST)
                use_selection = True
            except tk.TclError:
                # No selection, use entire document
                text_to_translate = self.text_area.get(1.0, tk.END).strip()
                use_selection = False
            
            if not text_to_translate:
                messagebox.showwarning("Uyarı", "Çevrilecek metin yok!")
                return
            
            # Show progress message
            self.statusbar.config(text="Çeviri yapılıyor...")
            self.root.update()
            
            # Perform translation in a thread to avoid blocking UI
            def do_translate():
                try:
                    translation = self.translator.translate(text_to_translate, dest=target_lang)
                    
                    # Update UI in main thread
                    self.root.after(0, lambda: self.show_translation(translation.text, use_selection))
                except Exception as e:
                    self.root.after(0, lambda: messagebox.showerror("Çeviri Hatası", 
                                                                    f"Çeviri sırasında hata: {str(e)}"))
                    self.root.after(0, lambda: self.statusbar.config(text="Hazır"))
            
            thread = threading.Thread(target=do_translate)
            thread.daemon = True
            thread.start()
            
        except Exception as e:
            messagebox.showerror("Hata", f"Çeviri hatası: {str(e)}")
            self.statusbar.config(text="Hazır")
            
    def show_translation(self, translated_text, replace_selection):
        """Show translation result"""
        # Create a new window to show translation
        trans_window = tk.Toplevel(self.root)
        trans_window.title("Çeviri Sonucu")
        trans_window.geometry("600x400")
        
        # Text area for translation
        trans_text = scrolledtext.ScrolledText(trans_window, wrap=tk.WORD, 
                                              font=(self.current_font_family, self.current_font_size))
        trans_text.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        trans_text.insert(1.0, translated_text)
        
        # Button frame
        btn_frame = tk.Frame(trans_window)
        btn_frame.pack(fill=tk.X, padx=10, pady=5)
        
        def replace_original():
            if replace_selection:
                try:
                    self.text_area.delete(tk.SEL_FIRST, tk.SEL_LAST)
                    self.text_area.insert(tk.INSERT, translated_text)
                except tk.TclError:
                    pass
            else:
                self.text_area.delete(1.0, tk.END)
                self.text_area.insert(1.0, translated_text)
            trans_window.destroy()
            self.statusbar.config(text="Çeviri tamamlandı")
        
        def copy_to_clipboard():
            self.root.clipboard_clear()
            self.root.clipboard_append(translated_text)
            messagebox.showinfo("Başarılı", "Çeviri panoya kopyalandı!")
        
        tk.Button(btn_frame, text="Orijinali Değiştir", 
                 command=replace_original).pack(side=tk.LEFT, padx=5)
        tk.Button(btn_frame, text="Panoya Kopyala", 
                 command=copy_to_clipboard).pack(side=tk.LEFT, padx=5)
        tk.Button(btn_frame, text="Kapat", 
                 command=trans_window.destroy).pack(side=tk.LEFT, padx=5)
        
        self.statusbar.config(text="Çeviri tamamlandı")
        
    def start_dictation(self):
        """Start voice dictation"""
        if not self.recognizer:
            messagebox.showerror("Hata", "Ses tanıma desteği yüklü değil. SpeechRecognition kütüphanesi gerekli.")
            return
            
        # Create dictation window
        dict_window = tk.Toplevel(self.root)
        dict_window.title("Sesle Metin Ekleme")
        dict_window.geometry("400x200")
        
        status_label = tk.Label(dict_window, text="Mikrofona konuşmaya başlayın...", 
                               font=("Arial", 12))
        status_label.pack(pady=20)
        
        progress = ttk.Progressbar(dict_window, mode='indeterminate')
        progress.pack(pady=10, padx=20, fill=tk.X)
        progress.start()
        
        result_label = tk.Label(dict_window, text="", wraplength=350)
        result_label.pack(pady=10)
        
        def recognize_speech():
            try:
                with sr.Microphone() as source:
                    # Adjust for ambient noise
                    self.recognizer.adjust_for_ambient_noise(source, duration=0.5)
                    
                    # Listen for speech
                    audio = self.recognizer.listen(source, timeout=5, phrase_time_limit=10)
                    
                    # Recognize speech using Google Speech Recognition
                    text = self.recognizer.recognize_google(audio, language='tr-TR')
                    
                    # Process voice commands
                    processed_text = self.process_voice_commands(text)
                    
                    # Insert text at cursor position
                    self.root.after(0, lambda: self.text_area.insert(tk.INSERT, processed_text + " "))
                    self.root.after(0, lambda: result_label.config(
                        text=f"Eklendi: {processed_text}", fg="green"))
                    self.root.after(0, lambda: progress.stop())
                    self.root.after(0, lambda: status_label.config(text="Başarılı!"))
                    
            except sr.WaitTimeoutError:
                self.root.after(0, lambda: result_label.config(
                    text="Zaman aşımı. Ses algılanamadı.", fg="red"))
                self.root.after(0, lambda: progress.stop())
                self.root.after(0, lambda: status_label.config(text="Hata!"))
            except sr.UnknownValueError:
                self.root.after(0, lambda: result_label.config(
                    text="Ses anlaşılamadı. Lütfen tekrar deneyin.", fg="red"))
                self.root.after(0, lambda: progress.stop())
                self.root.after(0, lambda: status_label.config(text="Hata!"))
            except Exception as e:
                self.root.after(0, lambda: result_label.config(
                    text=f"Hata: {str(e)}", fg="red"))
                self.root.after(0, lambda: progress.stop())
                self.root.after(0, lambda: status_label.config(text="Hata!"))
        
        # Start recognition in a separate thread
        thread = threading.Thread(target=recognize_speech)
        thread.daemon = True
        thread.start()
        
        tk.Button(dict_window, text="Kapat", 
                 command=dict_window.destroy).pack(pady=10)
        
    def process_voice_commands(self, text):
        """Process voice commands in text"""
        # Convert to lowercase for command detection
        text_lower = text.lower()
        
        # Handle "satır başı" command
        if "satır başı" in text_lower or "yeni satır" in text_lower:
            # Remove the command from text and add newline
            text = text_lower.replace("satır başı", "\n").replace("yeni satır", "\n")
            
        # Handle "nokta" command
        if text_lower.endswith("nokta"):
            text = text[:-5] + "."
            
        # Handle "virgül" command
        if "virgül" in text_lower:
            text = text.replace("virgül", ",")
            
        return text
        
    def add_line_marker(self):
        """Add visual marker to current line"""
        try:
            # Get current line
            current_line = self.text_area.index(tk.INSERT).split('.')[0]
            start = f"{current_line}.0"
            end = f"{current_line}.end"
            
            # Check if line already has marker
            tags = self.text_area.tag_names(start)
            if "marker" in tags:
                # Remove marker
                self.text_area.tag_remove("marker", start, end)
                self.statusbar.config(text="İşaretleyici kaldırıldı")
            else:
                # Add marker
                self.text_area.tag_add("marker", start, end)
                self.statusbar.config(text="Satıra işaretleyici eklendi")
        except Exception as e:
            messagebox.showerror("Hata", f"İşaretleyici eklenirken hata: {str(e)}")
            
    def undo(self):
        """Undo last action"""
        try:
            self.text_area.edit_undo()
        except tk.TclError:
            pass
            
    def redo(self):
        """Redo last undone action"""
        try:
            self.text_area.edit_redo()
        except tk.TclError:
            pass
            
    def cut(self):
        """Cut selected text"""
        try:
            self.text_area.event_generate("<<Cut>>")
        except tk.TclError:
            pass
            
    def copy(self):
        """Copy selected text"""
        try:
            self.text_area.event_generate("<<Copy>>")
        except tk.TclError:
            pass
            
    def paste(self):
        """Paste text from clipboard"""
        try:
            self.text_area.event_generate("<<Paste>>")
        except tk.TclError:
            pass
            
    def update_statusbar(self, event=None):
        """Update status bar with cursor position and character count"""
        # Get cursor position
        cursor_pos = self.text_area.index(tk.INSERT)
        line, col = cursor_pos.split('.')
        
        # Get character count
        content = self.text_area.get(1.0, tk.END)
        char_count = len(content) - 1  # Subtract 1 for the trailing newline
        word_count = len(content.split())
        
        self.statusbar.config(text=f"Satır: {line} | Sütun: {col} | "
                                  f"Karakter: {char_count} | Kelime: {word_count}")
        
    def show_about(self):
        """Show about dialog"""
        about_text = """
Gerekçeli Karar Programı
Versiyon 1.0

Bu program gerekçeli kararları düzenlemek, 
çevirmek ve çıktı almak amacıyla geliştirilmiştir.

Özellikler:
• Metin düzenleme ve formatlama
• Çoklu dil çevirisi
• Sesle metin ekleme
• PDF ve Word çıktısı

© 2024
        """
        messagebox.showinfo("Hakkında", about_text)


def main():
    """Main entry point"""
    root = tk.Tk()
    app = GerekceliKararProgram(root)
    root.mainloop()


if __name__ == "__main__":
    main()
